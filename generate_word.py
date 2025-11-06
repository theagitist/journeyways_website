#!/usr/bin/env python3
"""
Word Document Generator for JOURNEYWAYS Game Rules
Converts the rules.html page to a well-formatted Word document.
"""

import os
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def get_text_content(element):
    """Extract text content from an element, preserving card color spans."""
    text_parts = []
    for content in element.contents:
        if hasattr(content, 'name'):
            if content.name == 'span' and 'card-' in ' '.join(content.get('class', [])):
                # Preserve card color information
                classes = content.get('class', [])
                if 'card-green' in classes:
                    text_parts.append(('Green', content.get_text()))
                elif 'card-black' in classes:
                    text_parts.append(('Black', content.get_text()))
                elif 'card-blue' in classes:
                    text_parts.append(('Blue', content.get_text()))
                elif 'card-red' in classes:
                    text_parts.append(('Red', content.get_text()))
                elif 'card-purple' in classes:
                    text_parts.append(('Purple', content.get_text()))
                else:
                    text_parts.append(('', content.get_text()))
            else:
                text_parts.append(('', content.get_text() if hasattr(content, 'get_text') else str(content)))
        else:
            text_parts.append(('', str(content)))
    return text_parts

def add_formatted_text(paragraph, text_parts, default_color=None):
    """Add formatted text to a paragraph, preserving card colors."""
    for color_type, text in text_parts:
        if not text.strip():
            continue
        run = paragraph.add_run(text)
        if color_type == 'Green':
            run.font.color.rgb = RGBColor(16, 185, 129)  # #10b981
            run.bold = True
        elif color_type == 'Black':
            run.font.color.rgb = RGBColor(107, 114, 128)  # #6b7280
            run.bold = True
        elif color_type == 'Blue':
            run.font.color.rgb = RGBColor(59, 130, 246)  # #3b82f6
            run.bold = True
        elif color_type == 'Red':
            run.font.color.rgb = RGBColor(239, 68, 68)  # #ef4444
            run.bold = True
        elif color_type == 'Purple':
            run.font.color.rgb = RGBColor(139, 92, 246)  # #8b5cf6
            run.bold = True
        elif default_color:
            run.font.color.rgb = default_color

def generate_word_doc():
    """Generate Word document from the rules.html file."""
    
    # Define file paths
    html_file = '/var/www/www.journeywaysgame.com/rules.html'
    output_file = '/var/www/www.journeywaysgame.com/JOURNEYWAYS_Game_Rules.docx'
    
    # Check if HTML file exists
    if not os.path.exists(html_file):
        print(f"Error: HTML file not found at {html_file}")
        return False
    
    try:
        # Read and parse HTML
        print("Reading HTML file...")
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading('JOURNEYWAYS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(32)
        title_run.font.color.rgb = RGBColor(245, 158, 11)  # Yellow-400
        
        subtitle = doc.add_heading('Game Rules', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(24)
        subtitle_run.font.color.rgb = RGBColor(245, 158, 11)
        
        tagline = doc.add_paragraph('How to begin your journey of becoming')
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_run = tagline.runs[0]
        tagline_run.font.size = Pt(14)
        tagline_run.italic = True
        
        doc.add_paragraph()  # Spacing
        
        # Main heading
        main_heading = doc.add_heading('How to Play JOURNEYWAYS', 1)
        main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_heading_run = main_heading.runs[0]
        main_heading_run.font.color.rgb = RGBColor(245, 158, 11)
        
        # Introduction paragraph
        intro_text = "JOURNEYWAYS is not about winning or losing; it's about discovering, exploring, and becoming. These guidelines will help you create meaningful experiences, whether playing solo or with others."
        intro_para = doc.add_paragraph(intro_text)
        intro_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro_para_run = intro_para.runs[0]
        intro_para_run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
        
        # Find all rule sections
        rule_sections = soup.find_all('div', class_='rule-section')
        
        for section in rule_sections:
            # Get section title (h3)
            h3 = section.find('h3')
            if h3:
                section_title = h3.get_text(strip=True)
                heading = doc.add_heading(section_title, 2)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(245, 158, 11)
            
            # Process content based on section
            if 'Game Setup' in section_title:
                # Two columns: What You'll Need and Initial Setup
                h4_list = section.find_all('h4')
                for h4 in h4_list:
                    subheading = doc.add_heading(h4.get_text(strip=True), 3)
                    subheading_run = subheading.runs[0]
                    subheading_run.font.color.rgb = RGBColor(55, 65, 81)  # Gray-700
                    
                    # Find the list after this h4
                    parent = h4.find_parent()
                    if parent:
                        ul = parent.find('ul')
                        ol = parent.find('ol')
                        
                        if ul:
                            for li in ul.find_all('li', recursive=False):
                                text = li.get_text(strip=True)
                                # Remove bullet if present
                                if text.startswith('•'):
                                    text = text[1:].strip()
                                para = doc.add_paragraph(text, style='List Bullet')
                                para_run = para.runs[0]
                                para_run.font.size = Pt(11)
                        
                        if ol:
                            for li in ol.find_all('li', recursive=False):
                                text = li.get_text(strip=True)
                                # Remove number if present
                                text = text.lstrip('0123456789. ').strip()
                                para = doc.add_paragraph(text, style='List Number')
                                para_run = para.runs[0]
                                para_run.font.size = Pt(11)
            
            elif 'Basic Gameplay' in section_title:
                # Turn Structure section
                h4 = section.find('h4')
                if h4 and 'Turn Structure' in h4.get_text():
                    subheading = doc.add_heading(h4.get_text(strip=True), 3)
                    subheading_run = subheading.runs[0]
                    subheading_run.font.color.rgb = RGBColor(55, 65, 81)
                    
                    # Get the paragraph before the grid
                    p = h4.find_next_sibling('p')
                    if p:
                        para = doc.add_paragraph()
                        text_parts = get_text_content(p)
                        add_formatted_text(para, text_parts)
                        para_run = para.runs[0]
                        para_run.font.size = Pt(11)
                    
                    # Process the three phases
                    phase_divs = section.find_all('div', class_='bg-gray-700')
                    for phase_div in phase_divs:
                        h5 = phase_div.find('h5')
                        if h5:
                            phase_heading = doc.add_heading(h5.get_text(strip=True), 4)
                            phase_heading_run = phase_heading.runs[0]
                            phase_heading_run.font.color.rgb = RGBColor(245, 158, 11)
                        
                        p = phase_div.find('p')
                        if p:
                            para = doc.add_paragraph()
                            text_parts = get_text_content(p)
                            add_formatted_text(para, text_parts)
                            para_run = para.runs[0]
                            para_run.font.size = Pt(10)
            
            elif 'Ending the game' in section_title:
                # Paragraph content
                paragraphs = section.find_all('p', class_='text-gray-300')
                for p in paragraphs:
                    para = doc.add_paragraph()
                    text_parts = get_text_content(p)
                    add_formatted_text(para, text_parts)
                    for run in para.runs:
                        run.font.size = Pt(11)
            
            elif 'Writing your journal' in section_title:
                paragraphs = section.find_all('p', class_='text-gray-300')
                for p in paragraphs:
                    para = doc.add_paragraph()
                    text_parts = get_text_content(p)
                    add_formatted_text(para, text_parts)
                    for run in para.runs:
                        run.font.size = Pt(11)
            
            elif 'Solo vs Group Play' in section_title:
                h4_list = section.find_all('h4')
                for h4 in h4_list:
                    subheading = doc.add_heading(h4.get_text(strip=True), 3)
                    subheading_run = subheading.runs[0]
                    if 'Solo' in h4.get_text():
                        subheading_run.font.color.rgb = RGBColor(16, 185, 129)  # Green
                    else:
                        subheading_run.font.color.rgb = RGBColor(59, 130, 246)  # Blue
                    
                    # Find the list after this h4
                    parent = h4.find_parent()
                    if parent:
                        ul = parent.find('ul')
                        if ul:
                            for li in ul.find_all('li', recursive=False):
                                para = doc.add_paragraph()
                                text_parts = get_text_content(li)
                                add_formatted_text(para, text_parts)
                                para.style = 'List Bullet'
                                for run in para.runs:
                                    run.font.size = Pt(11)
            
            elif 'Advanced Concepts' in section_title:
                h4_list = section.find_all('h4')
                for h4 in h4_list:
                    subheading = doc.add_heading(h4.get_text(strip=True), 3)
                    subheading_run = subheading.runs[0]
                    subheading_run.font.color.rgb = RGBColor(55, 65, 81)
                    
                    # Find the paragraph after this h4
                    p = h4.find_next_sibling('p')
                    if p:
                        para = doc.add_paragraph()
                        text_parts = get_text_content(p)
                        add_formatted_text(para, text_parts)
                        for run in para.runs:
                            run.font.size = Pt(11)
            
            elif 'Tips for Meaningful Play' in section_title:
                h4_list = section.find_all('h4')
                for h4 in h4_list:
                    subheading = doc.add_heading(h4.get_text(strip=True), 3)
                    subheading_run = subheading.runs[0]
                    subheading_run.font.color.rgb = RGBColor(55, 65, 81)
                    
                    # Find the list after this h4
                    parent = h4.find_parent()
                    if parent:
                        ul = parent.find('ul')
                        if ul:
                            for li in ul.find_all('li', recursive=False):
                                text = li.get_text(strip=True)
                                if text.startswith('•'):
                                    text = text[1:].strip()
                                para = doc.add_paragraph(text, style='List Bullet')
                                para_run = para.runs[0]
                                para_run.font.size = Pt(11)
            
            doc.add_paragraph()  # Spacing between sections
        
        # Add footer note
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run('© 2025 JOURNEYWAYS. A board game about becoming.')
        footer_run.font.size = Pt(10)
        footer_run.italic = True
        footer_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Save document
        print("Generating Word document...")
        doc.save(output_file)
        
        print(f"Word document successfully generated: {output_file}")
        return True
        
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = generate_word_doc()
    sys.exit(0 if success else 1)

